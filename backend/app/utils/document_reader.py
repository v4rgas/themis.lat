"""
Utility functions for document type detection and local text extraction.
"""
import io
from typing import Dict, Optional, Any
import puremagic
from docx import Document


def detect_file_type(file_content: bytes) -> str:
    """
    Detect the MIME type of a file from its content.
    
    Args:
        file_content: Raw file content as bytes
        
    Returns:
        MIME type string (e.g., "application/pdf", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    """
    try:
        # puremagic.magic_string returns a list of matches: [(extension, mime_type, description, confidence), ...]
        matches = puremagic.magic_string(file_content)
        if matches and len(matches) > 0:
            # Get the first match (highest confidence)
            match = matches[0]
            if len(match) >= 2:
                mime_type = match[1]  # MIME type is the second element
                if mime_type:
                    return mime_type
    except Exception:
        pass
    
    # Fallback: try to detect from magic bytes
    try:
        if file_content.startswith(b'%PDF'):
            return "application/pdf"
        elif file_content.startswith(b'PK\x03\x04'):
            # Could be DOCX, XLSX, PPTX, etc. - check more carefully
            # DOCX files have specific structure
            if b'word/' in file_content[:1024]:
                return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    except Exception:
        pass
    
    raise Exception("Failed to detect file type")


def get_file_extension_from_mime(mime_type: str) -> str:
    """
    Map MIME type to file extension.
    
    Args:
        mime_type: MIME type string
        
    Returns:
        File extension with leading dot (e.g., ".pdf", ".docx")
    """
    mime_to_ext = {
        "application/pdf": ".pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
        "application/msword": ".doc",
        "image/png": ".png",
        "image/jpeg": ".jpg",
        "image/jpg": ".jpg",
        "image/gif": ".gif",
        "image/tiff": ".tiff",
        "image/bmp": ".bmp",
    }
    
    extension = mime_to_ext.get(mime_type)
    if extension:
        return extension
    
    # Default fallback
    return ".pdf"


def extract_text_from_docx(file_content: bytes) -> Dict[str, Any]:
    """
    Extract all text from a DOCX file.
    
    Args:
        file_content: Raw DOCX file content as bytes
        
    Returns:
        dict with keys: {text: str, success: bool, error?: str}
    """
    try:
        # Create a file-like object from bytes
        docx_file = io.BytesIO(file_content)
        doc = Document(docx_file)
        
        # Extract text from all paragraphs
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        combined_text = "\n\n".join(text_parts)
        
        return {
            "text": combined_text,
            "success": True
        }
    except Exception as e:
        return {
            "text": "",
            "success": False,
            "error": f"Failed to extract text from DOCX: {str(e)}"
        }


def extract_text_locally(file_content: bytes, file_type: str) -> Dict[str, Any]:
    """
    Extract text from a document locally if supported.
    
    Args:
        file_content: Raw file content as bytes
        file_type: MIME type of the file
        
    Returns:
        dict with keys: {text: str, success: bool, error?: str}
    """
    # Only DOCX is supported for local extraction
    if file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return extract_text_from_docx(file_content)
    
    # For other file types (PDF, images), return success=False to indicate
    # that local extraction is not supported/attempted
    return {
        "text": "",
        "success": False,
        "error": f"Local extraction not supported for {file_type}"
    }

